from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

from app.core.errors import AppError


@dataclass(frozen=True)
class TextExtractionResult:
    raw_text: str
    extraction_status: str
    extraction_method: str
    warnings: list[str] = field(default_factory=list)


TEXT_FILE_TYPES = {"markdown", "text"}


def extract_resume_text(
    filename: str, file_type: str, content_type: str | None, content: bytes
) -> TextExtractionResult:
    if file_type in TEXT_FILE_TYPES:
        return _extract_utf8_text(filename, file_type, content)
    if file_type == "pdf":
        return _extract_pdf_text(filename, content)
    if file_type == "docx":
        return _extract_docx_text(filename, content)
    raise AppError(
        code="unsupported_resume_file_type",
        message="Supported resume file types are PDF, DOCX, Markdown, and text.",
        status_code=400,
        details={
            "filename": filename,
            "content_type": content_type,
            "file_type": file_type,
        },
    )


def _extract_utf8_text(
    filename: str, file_type: str, content: bytes
) -> TextExtractionResult:
    try:
        raw_text = content.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise AppError(
            code="resume_text_decode_failed",
            message="Resume text file must be valid UTF-8.",
            status_code=400,
            details={"filename": filename, "file_type": file_type},
        ) from exc

    normalized_text = raw_text.strip()
    if not normalized_text:
        raise AppError(
            code="resume_text_empty_after_extraction",
            message="Resume text extraction produced empty text.",
            status_code=400,
            details={"filename": filename, "file_type": file_type},
        )

    suffix = Path(filename).suffix.lower().lstrip(".") or file_type
    return TextExtractionResult(
        raw_text=normalized_text,
        extraction_status="extracted",
        extraction_method=f"utf8_{suffix}_decode",
        warnings=[],
    )


def _extract_pdf_text(filename: str, content: bytes) -> TextExtractionResult:
    try:
        import fitz

        with fitz.open(stream=content, filetype="pdf") as document:
            page_text = [page.get_text("text").strip() for page in document]
    except Exception as exc:
        raise AppError(
            code="resume_pdf_extract_failed",
            message="Unable to extract text from PDF resume.",
            status_code=400,
            details={"filename": filename, "file_type": "pdf"},
        ) from exc

    return _build_extracted_result(
        filename=filename,
        file_type="pdf",
        raw_text="\n\n".join(text for text in page_text if text),
        method="pymupdf_text",
    )


def _extract_docx_text(filename: str, content: bytes) -> TextExtractionResult:
    try:
        from docx import Document

        document = Document(BytesIO(content))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    except Exception as exc:
        raise AppError(
            code="resume_docx_extract_failed",
            message="Unable to extract text from DOCX resume.",
            status_code=400,
            details={"filename": filename, "file_type": "docx"},
        ) from exc

    return _build_extracted_result(
        filename=filename,
        file_type="docx",
        raw_text="\n".join(part for part in parts if part),
        method="python_docx_text",
    )


def _build_extracted_result(
    *, filename: str, file_type: str, raw_text: str, method: str
) -> TextExtractionResult:
    normalized_text = raw_text.strip()
    if not normalized_text:
        raise AppError(
            code="resume_text_empty_after_extraction",
            message="Resume text extraction produced empty text.",
            status_code=400,
            details={"filename": filename, "file_type": file_type},
        )
    return TextExtractionResult(
        raw_text=normalized_text,
        extraction_status="extracted",
        extraction_method=method,
        warnings=[],
    )
